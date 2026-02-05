"""
Text extraction from uploaded files (PDF, DOCX, TXT).

Converts file uploads to RawUpload model with extracted text content.
"""

import os
from datetime import datetime

from invitation_triage.models.upload import RawUpload


def extract_text_from_pdf(file_path: str) -> tuple[str, dict]:
    """
    Extract text from PDF using pypdf.

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    metadata = {
        "page_count": len(reader.pages),
    }

    return "\n\n".join(text_parts), metadata


def extract_text_from_docx(file_path: str) -> tuple[str, dict]:
    """
    Extract text from Word document using python-docx.

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    metadata = {
        "paragraph_count": len(doc.paragraphs),
    }

    return "\n\n".join(text_parts), metadata


def extract_text_from_txt(file_path: str) -> tuple[str, dict]:
    """
    Read plain text file.

    Returns:
        Tuple of (file_content, metadata_dict)
    """
    with open(file_path, encoding="utf-8") as f:
        text = f.read()

    metadata = {
        "line_count": text.count("\n") + 1,
    }

    return text, metadata


def extract_text_from_file(file_path: str, file_type: str | None = None) -> RawUpload:
    """
    Main function - extract text from file and create RawUpload.

    Args:
        file_path: Path to the file to process
        file_type: Optional file type override ("pdf", "docx", "txt").
                   If not provided, will be inferred from file extension.

    Returns:
        RawUpload with extracted text and metadata

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # Infer file type from extension if not provided
    if file_type is None:
        ext = os.path.splitext(file_path)[1].lower()
        type_mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "docx",
            ".txt": "txt",
        }
        file_type = type_mapping.get(ext)
        if file_type is None:
            raise ValueError(
                f"Cannot infer file type from extension '{ext}'. "
                f"Supported extensions: {list(type_mapping.keys())}"
            )

    # Extract text based on file type
    if file_type == "pdf":
        text, metadata = extract_text_from_pdf(file_path)
    elif file_type == "docx":
        text, metadata = extract_text_from_docx(file_path)
    elif file_type == "txt":
        text, metadata = extract_text_from_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {file_type}. "
            f"Supported types: pdf, docx, txt"
        )

    # Generate upload ID
    filename = os.path.basename(file_path)
    upload_id = RawUpload._generate_upload_id(text, filename)

    return RawUpload(
        upload_id=upload_id,
        filename=filename,
        source_type=file_type,
        raw_text=text,
        upload_timestamp=datetime.now(),
        file_size=os.path.getsize(file_path),
        metadata=metadata,
    )
