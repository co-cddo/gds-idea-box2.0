"""
Text extraction from documented files (PDF, DOCX, TXT).

Converts file documents to RawDocument model with extracted text content.
"""

import os
import tempfile
from datetime import datetime

from invitation_triage.models.document import RawDocument


def extract_text_from_pdf(file_path: str) -> tuple[str, dict]:
    """
    Extract text from PDF using pypdf.

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    metadata = {
        "page_count": len(reader.pages),
    }

    return "\n\n".join(text_parts), metadata


def extract_text_from_docx(file_path: str) -> tuple[str, dict]:
    """
    Extract text from Word document using python-docx.

    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    metadata = {
        "paragraph_count": len(doc.paragraphs),
    }

    return "\n\n".join(text_parts), metadata


def extract_text_from_txt(file_path: str) -> tuple[str, dict]:
    """
    Read plain text file.

    Returns:
        Tuple of (file_content, metadata_dict)
    """
    with open(file_path, encoding="utf-8") as f:
        text = f.read()

    metadata = {
        "line_count": text.count("\n") + 1,
    }

    return text, metadata


def extract_text_from_file(
    file_path: str | None = None,
    file_type: str | None = None,
    binary_data: bytes | None = None,
    filename: str | None = None,
) -> RawDocument:
    """
    Extract text from file (path or binary data) and create RawDocument.

    Args:
        file_path: Path to file (for direct file uploads)
        file_type: File type override ("pdf", "docx", "txt")
        binary_data: Raw bytes (for email attachments)
        filename: Original filename (required if using binary_data)

    Returns:
        RawDocument with extracted text and metadata

    Raises:
        ValueError: If file type is not supported or invalid parameters
        FileNotFoundError: If file_path doesn't exist
    """
    # Validate inputs
    if file_path is None and binary_data is None:
        raise ValueError("Must provide either file_path or binary_data")
    if binary_data is not None and filename is None:
        raise ValueError("filename required when using binary_data")

    # Handle binary data by writing to temp file
    temp_file_path = None
    if binary_data is not None:
        # Infer file type from filename if not provided
        if file_type is None and filename:
            ext = os.path.splitext(filename)[1].lower()
            type_mapping = {
                ".pdf": "pdf",
                ".docx": "docx",
                ".doc": "docx",
                ".txt": "txt",
            }
            file_type = type_mapping.get(ext)

        if file_type is None:
            raise ValueError(
                "Cannot determine file type. Provide file_type parameter "
                "or use standard file extension in filename"
            )

        # Write binary data to temp file
        with tempfile.NamedTemporaryFile(
            delete=False, suffix=f".{file_type}"
        ) as tmp:
            tmp.write(binary_data)
            temp_file_path = tmp.name

        # Use temp file as file_path for processing
        file_path = temp_file_path
        file_size = len(binary_data)
    else:
        # Using file_path directly
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        file_size = os.path.getsize(file_path)
        if filename is None:
            filename = os.path.basename(file_path)

    try:
        # Infer file type from extension if not provided
        if file_type is None:
            ext = os.path.splitext(file_path)[1].lower()
            type_mapping = {
                ".pdf": "pdf",
                ".docx": "docx",
                ".doc": "docx",
                ".txt": "txt",
            }
            file_type = type_mapping.get(ext)
            if file_type is None:
                raise ValueError(
                    f"Cannot infer file type from extension '{ext}'. "
                    f"Supported extensions: {list(type_mapping.keys())}"
                )

        # Extract text based on file type
        if file_type == "pdf":
            text, metadata = extract_text_from_pdf(file_path)
        elif file_type == "docx":
            text, metadata = extract_text_from_docx(file_path)
        elif file_type == "txt":
            text, metadata = extract_text_from_txt(file_path)
        else:
            raise ValueError(
                f"Unsupported file type: {file_type}. Supported types: pdf, docx, txt"
            )

        # Generate document ID
        document_id = RawDocument._generate_document_id(text, filename)

        return RawDocument(
            document_id=document_id,
            filename=filename,
            source_type=file_type,
            raw_text=text,
            document_timestamp=datetime.now(),
            file_size=file_size,
            metadata=metadata,
        )

    finally:
        # Clean up temp file if we created one
        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
